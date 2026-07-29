from pathlib import Path
from docx import Document
from docx.shared import Pt
import config

SALIDA_DIR = Path("cvs_generados")


def _agregar_titulo_seccion(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)


def generar_cv(
    vacante: dict,
    habilidades_cubiertas: list[str],
    perfil_override: str | None = None,
    habilidades_orden_override: list[str] | None = None,
) -> Path:
    """
    Construye el CV completo en Word. Los datos fijos (estudios, experiencia)
    salen de config.DATOS_CV. El perfil y el orden de habilidades se adaptan
    a las palabras clave que sí coinciden con la vacante — o, si el usuario
    ya editó el CV a mano desde el dashboard, se usan esos textos tal cual
    (perfil_override / habilidades_orden_override).
    """
    datos = config.DATOS_CV
    doc = Document()

    doc.add_heading(datos["nombre"], level=1)
    doc.add_paragraph(f"{datos['telefono']} | {datos['email']}")

    _agregar_titulo_seccion(doc, "Perfil")
    if perfil_override:
        perfil = perfil_override
    else:
        destacadas = ", ".join(habilidades_cubiertas[:5]) if habilidades_cubiertas else "automatización de procesos y desarrollo en Python"
        perfil = (
            f"Técnico en Programación con experiencia en automatización de procesos, desarrollo de "
            f"herramientas en Python y gestión de bases de datos. Cuento con experiencia real en entorno "
            f"corporativo (Marsh McLennan). Para el cargo de {vacante['cargo']} en {vacante['empresa']}, "
            f"destaco mi manejo de {destacadas}. Actualmente cursando Ingeniería de Software, enfocado en "
            f"buenas prácticas de desarrollo y mejora continua."
        )
    doc.add_paragraph(perfil)

    _agregar_titulo_seccion(doc, "Habilidades Técnicas")
    if habilidades_orden_override:
        doc.add_paragraph(", ".join(habilidades_orden_override))
    else:
        for categoria, items in datos["categorias_habilidades"].items():
            # las habilidades que coinciden con la vacante van primero
            items_ordenados = sorted(
                items, key=lambda h: h.lower() not in [x.lower() for x in habilidades_cubiertas]
            )
            doc.add_paragraph(f"{categoria}: {', '.join(items_ordenados)}")

    _agregar_titulo_seccion(doc, "Estudios")
    for institucion, titulo, periodo in datos["estudios"]:
        doc.add_paragraph(f"{institucion} — {titulo} ({periodo})")

    _agregar_titulo_seccion(doc, "Experiencia Laboral")
    for exp in datos["experiencia"]:
        doc.add_paragraph(f"{exp['empresa']} — {exp['rol']} ({exp['periodo']})", style="List Bullet")
        for logro in exp["logros"]:
            doc.add_paragraph(logro, style="List Bullet 2")

    _agregar_titulo_seccion(doc, "Información Adicional")
    doc.add_paragraph(f"Idioma: {datos['idiomas']}")

    SALIDA_DIR.mkdir(exist_ok=True)
    nombre_archivo = f"CV_{vacante['empresa']}_{vacante['cargo']}"
    nombre_archivo = "".join(c for c in nombre_archivo if c.isalnum() or c in " _-").replace(" ", "_")
    ruta = SALIDA_DIR / f"{nombre_archivo}.docx"
    doc.save(ruta)
    return ruta
